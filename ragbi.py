import os
import yaml
import locale
import warnings
import pytesseract
import re
import hashlib
import tempfile
import time
from datetime import datetime, timedelta
from langdetect import detect, DetectorFactory
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_google_genai import ChatGoogleGenerativeAI
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# Suppress warnings and set encoding
warnings.filterwarnings("ignore")
locale.getpreferredencoding = lambda: 'UTF-8'

# Fix for langdetect consistency
DetectorFactory.seed = 0

def load_gemini_api_key():
    """Load Gemini API key from YAML configuration file"""
    try:
        with open('gemini_api.yml', 'r') as file:
            api_creds = yaml.safe_load(file)
            os.environ['GOOGLE_API_KEY'] = api_creds['api']['api_key']
    except Exception as e:
        print(f"Error loading Gemini API key: {str(e)}")
        raise

# Load API key on import
load_gemini_api_key()

class GeminiLLM:
    def __init__(self, model="gemini-1.5-flash", rate_limit=60):
        self.last_call = datetime.now() - timedelta(seconds=rate_limit)
        self.gemini_model = model
        self.rate_limit = rate_limit

    def generate(self, prompt, is_bengali):
        """Generate response from Gemini model with rate limiting"""
        try:
            # Enforce rate limiting
            now = datetime.now()
            elapsed = (now - self.last_call).total_seconds()
            if elapsed < self.rate_limit:
                time.sleep(self.rate_limit - elapsed)
            
            # Add Bengali instruction if needed
            if is_bengali:
                prompt = f"এই প্রশ্নের উত্তর বাংলায় দিন:\n{prompt}"
            
            # Initialize Gemini model
            gemini = ChatGoogleGenerativeAI(
                model=self.gemini_model,
                temperature=0.2,
                convert_system_message_to_human=True
            )
            
            # Generate response
            response = gemini.invoke(prompt)
            self.last_call = datetime.now()
            return response.content
        
        except Exception as e:
            error_msg = str(e)
            if "quota" in error_msg.lower() or "429" in error_msg:
                return "API quota exceeded. Please try again later."
            elif "location is not supported" in error_msg:
                self.gemini_model = "gemini-1.0-pro"
                return self.generate(prompt, is_bengali)  # Retry with different model
            else:
                return f"Error generating response: {str(e)}"

def extract_text_with_ocr(pdf_path):
    """Extract Bengali text from PDF using OCR"""
    from pdf2image import convert_from_path
    try:
        # Convert PDF to images
        images = convert_from_path(
            pdf_path, 
            dpi=300, 
            poppler_path="/usr/bin",
            thread_count=4
        )
        
        text = ""
        for image in images:
            # OCR with Bengali language
            page_text = pytesseract.image_to_string(
                image.convert('L'),  # Convert to grayscale
                lang='ben',
                config='--psm 6 --oem 3'  # Page segmentation mode 6 (single uniform block)
            )
            # Clean Bengali text
            page_text = re.sub(r'[^\u0980-\u09FF\s.,!?;:]+', '', page_text)
            text += f"{page_text}\n\n"
        return text.strip()
    
    except Exception as e:
        return f"OCR Error: {str(e)}"

def get_file_hash(file_content):
    """Generate MD5 hash for file content"""
    return hashlib.md5(file_content).hexdigest()

def load_document(file_path, file_extension, file_name):
    """Load document content based on file type"""
    try:
        if file_extension == ".pdf":
            # Process PDF files
            doc = fitz.open(file_path)
            text = ""
            for i in range(len(doc)):
                page = doc.load_page(i)
                text += page.get_text() + "\n\n"
            
            # Check for Bengali characters and use OCR if needed
            if any('\u0980' <= char <= '\u09FF' for char in text):
                ocr_text = extract_text_with_ocr(file_path)
                if ocr_text and not ocr_text.startswith("OCR Error"):
                    return [Document(
                        page_content=ocr_text,
                        metadata={"source": file_name, "pages": f"1-{len(doc)}"}
                    )]
            
            return [Document(
                page_content=text,
                metadata={"source": file_name, "pages": f"1-{len(doc)}"}
            )]
        
        elif file_extension == ".docx":
            # Process DOCX files
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return [Document(
                page_content=text,
                metadata={"source": file_name, "pages": "all"}
            )]
        
        elif file_extension == ".txt":
            # Process TXT files
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return [Document(
                page_content=text,
                metadata={"source": file_name, "pages": "all"}
            )]
        
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    except Exception as e:
        print(f"Error loading document {file_name}: {str(e)}")
        return []

def detect_language(text):
    """Detect if text is Bengali or English"""
    try:
        lang = detect(text)
        return "bn" if lang in ["bn", "as", "or"] else "en"
    except:
        return "en"

def format_sources(docs):
    """Format document sources for display"""
    sources = []
    for doc in docs:
        source = doc.metadata.get("source", "Unknown")
        pages = doc.metadata.get("pages", "N/A")
        content = doc.page_content.strip()
        content = content[:150] + "..." if len(content) > 150 else content
        sources.append({
            "Source File": source,
            "Page(s)": pages,
            "Content Excerpt": content
        })
    return sources

def build_vectorstore(documents):
    """Build Chroma vector store from documents"""
    embedding_model = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )
    return Chroma.from_documents(
        documents=documents,
        embedding=embedding_model,
        persist_directory="./chroma_db"  # Optional persistence
    )

def process_documents(file_data_list):
    """Process uploaded files and return vector store"""
    all_docs = []
    processed_files = {}
    
    for file_name, file_content in file_data_list:
        file_hash = get_file_hash(file_content)
        if file_hash in processed_files:
            continue  # Skip already processed files
            
        _, file_extension = os.path.splitext(file_name)
        file_extension = file_extension.lower()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
        
        docs = load_document(tmp_path, file_extension, file_name)
        os.unlink(tmp_path)
        
        if docs:
            # Split document into chunks
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", "।", "। ", " ", ""]  # Bengali-aware splitting
            )
            chunks = splitter.split_documents(docs)
            all_docs.extend(chunks)
            processed_files[file_hash] = file_name
    
    # Create vector store if documents were processed
    if all_docs:
        return build_vectorstore(all_docs)
    return None

def get_ragbi_response(query, vectorstore, llm):
    """Get RAG response for a query"""
    if not vectorstore:
        return "Vector store not initialized. Please upload documents first.", []
    
    # Retrieve relevant documents
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    docs = retriever.get_relevant_documents(query)
    
    # Detect query language
    lang = detect_language(query)
    is_bengali = lang == "bn"
    
    # Format context with sources
    context = "\n\n".join([
        f"Source: {doc.metadata.get('source', 'Unknown')} (Pages: {doc.metadata.get('pages', 'N/A')})\n{doc.page_content}"
        for doc in docs
    ])
    
    # Create language-specific prompt
    if is_bengali:
        prompt = f"প্রসঙ্গ: {context}\nপ্রশ্ন: {query}\nউত্তর: "
    else:
        prompt = f"Context: {context}\nQuestion: {query}\nAnswer:"
    
    # Generate response
    answer = llm.generate(prompt, is_bengali)
    
    # Format sources
    sources = format_sources(docs)
    
    return answer, sources